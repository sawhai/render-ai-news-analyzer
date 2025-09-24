#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Aug 18 19:30:09 2025

@author: ha
"""
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
COMPLETELY FIXED FastAPI Backend for Kuwait Banking News Analyzer
This version resolves ALL identified issues:
1. Content transfer to Word documents
2. WebSocket connection stability
3. Frontend data synchronization
4. Report generation accuracy
"""
import tempfile
import shutil
from pathlib import Path
import os
import asyncio
import json
import uuid
import zipfile
import io
import subprocess
import sys
import tempfile
from datetime import datetime
from typing import Dict, List, Optional
from pathlib import Path
import re
from zoneinfo import ZoneInfo
from dotenv import load_dotenv
from openai import OpenAI
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from collections import defaultdict

KW = ZoneInfo("Asia/Kuwait")
load_dotenv()
openai_client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# CRITICAL FIX: Handle asyncio loop compatibility
import nest_asyncio

def fix_asyncio_loop():
    """Fix asyncio loop compatibility with nest_asyncio and uvloop"""
    try:
        nest_asyncio.apply()
        print("âœ… nest_asyncio patch applied successfully")
    except Exception as e:
        print(f"âš ï¸ nest_asyncio patch warning: {e}")
        # Set default event loop policy to avoid uvloop conflicts
        asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())
        print("âœ… Set default event loop policy")

# Apply the fix immediately
fix_asyncio_loop()

from fastapi import FastAPI, WebSocket, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
import uvicorn

# Add Word document generation capabilities
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from PIL import Image

# Import bank configurations
try:
    from bank_patterns import BANK_CONFIGS, get_bank_info
    print("✅ Successfully imported bank_patterns")
except ImportError as e:
    print(f"⚠️ Warning: Could not import bank_patterns: {e}")

    
    # Fallback bank configuration
    BANK_CONFIGS = {
        'gulf_bank': {'english_name': 'Gulf Bank', 'arabic_name': 'بنك الخليج'},
        'nbk': {'english_name': 'National Bank of Kuwait', 'arabic_name': 'بنك الكويت الوطني'},
        'kfh': {'english_name': 'Kuwait Finance House', 'arabic_name': 'بيت التمويل الكويتي'},
        'boubyan_bank': {'english_name': 'Boubyan Bank', 'arabic_name': 'بنك بوبيان'},
        'cbk': {'english_name': 'Commercial Bank of Kuwait', 'arabic_name': 'البنك التجاري الكويتي'},
        'burgan_bank': {'english_name': 'Burgan Bank', 'arabic_name': 'بنك برقان'},
        'kib': {'english_name': 'Kuwait International Bank', 'arabic_name': 'بنك الكويت الدولي'},
        'abk': {'english_name': 'Al Ahli Bank of Kuwait', 'arabic_name': 'بنك الاهلي الكويتي'},
        'warba_bank': {'english_name': 'Warba Bank', 'arabic_name': 'بنك وربة'}
    }
    
    def get_bank_info(bank_code):
        return BANK_CONFIGS.get(bank_code, {'english_name': bank_code, 'arabic_name': bank_code})

def setup_cloud_environment():
    """Setup environment for cloud deployment"""
    data_dir = os.getenv("DATA_DIR", "/tmp/bank_news_data")
    
    # Check available disk space
    import shutil
    total, used, free = shutil.disk_usage(data_dir)
    free_mb = free / (1024 * 1024)
    print(f"Available storage: {free_mb:.1f}MB")
    
    if free_mb < 100:  # Less than 100MB free
        print("WARNING: Low disk space available")
    
    # Create necessary directories
    data_dir = os.getenv("DATA_DIR", "/tmp/bank_news_data")
    os.makedirs(data_dir, exist_ok=True)
    os.makedirs(os.path.join(data_dir, "alrai"), exist_ok=True)
    os.makedirs(os.path.join(data_dir, "AlJarida"), exist_ok=True)
    os.makedirs(os.path.join(data_dir, "KuwaitTimes"), exist_ok=True)
    os.makedirs("Reports", exist_ok=True)
    os.makedirs("Charts", exist_ok=True)
    
    print(f"Cloud environment setup complete. Data directory: {data_dir}")

app = FastAPI(title="Kuwait Banking News Analyzer API", version="2.2.0")

# Cloud deployment configuration
cors_origins = [
    "http://localhost:3000", 
    "http://127.0.0.1:3000"
]

# Add production origins if in cloud environment
if os.getenv('RENDER'):
    cors_origins.extend([
        "https://your-frontend-domain.com",  # Replace with actual domain when you have it
        "https://*.render.com"
    ])

# Enable CORS for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

async def ensure_playwright_browser():
    """Ensure Playwright browser is available in cloud environment"""
    try:
        from playwright.async_api import async_playwright
        async with async_playwright() as p:
            # Test if browser is available
            browser = await p.chromium.launch(headless=True)
            await browser.close()
            print("Playwright browser check successful")
    except Exception as e:
        print(f"Playwright browser not available: {e}")
        # You might want to install it here or handle the error
        raise e

# Task storage - In production, use Redis or database
active_tasks: Dict[str, dict] = {}

class AnalysisRequest(BaseModel):
    selected_banks: List[str]
    selected_newspapers: List[str]
    selected_model: str = "gpt-4o"  # Add this field with default

# Newspaper script mapping
NEWSPAPER_SCRIPTS = {
    'alrai_multibank': 'Al-Rai Media',
    'aljarida_multibank': 'Al-Jarida',
    'alqabas_multibank': 'Al-Qabas',
    'alnahar_multibank': 'Al-Nahar',
    'kwttimes_multibank': 'Kuwait Times',
    'arabtimes_multibank': 'Arab Times',
    'alwasat_multibank': 'Al-Wasat',
    'alanbaa_multibanks': 'Al-Anbaa',
    'alseyassah_multibank': 'Al-Seyassah' 
}

# VALIDATION_CONFIG = {
#     'enabled': True,
#     'llm_validation_enabled': True,  # Add this
#     'llm_model': 'gpt-4o-mini',      # Add this
#     'debug_validation': True
# }

VALIDATION_CONFIG = {
    'enabled': True,
    'llm_validation_enabled': True,
    'llm_model': 'gpt-4o-mini',
    'debug_validation': True,
    'min_confidence_threshold': 0.5,
    'require_strict_criteria': True
}

# ==================================================================================
# FIXED MASTER ORCHESTRATOR INTEGRATION CLASSES
# ==================================================================================

class MultiBankAnalysisTracker:
    """FIXED Track the analysis status of each bank-newspaper combination"""
    def __init__(self):
        self.results = {}  # Structure: {bank_code: {newspaper: result}}
        self.newspaper_stats = {}  # Track newspaper-level stats
        self.bank_stats = {}  # Track bank-level stats
        self.total_combinations = 0
        self.completed_combinations = 0
        
    def add_result(self, bank_code, newspaper_script, newspaper_name, result):
        """Add analysis result for a bank-newspaper combination"""
        if bank_code not in self.results:
            self.results[bank_code] = {}
        
        self.results[bank_code][newspaper_script] = {
            'newspaper_name': newspaper_name,
            'result': result,
            'success': result is not None,
            'content_found': result.get('bank_found', False) if result else False
        }
        
        self.completed_combinations += 1
        
        # Update newspaper stats
        if newspaper_script not in self.newspaper_stats:
            self.newspaper_stats[newspaper_script] = {
                'name': newspaper_name,
                'banks_analyzed': 0,
                'banks_with_content': 0,
                'total_pages_found': 0
            }
        
        self.newspaper_stats[newspaper_script]['banks_analyzed'] += 1
        if result and result.get('bank_found'):
            self.newspaper_stats[newspaper_script]['banks_with_content'] += 1
            self.newspaper_stats[newspaper_script]['total_pages_found'] += len(result.get('relevant_pages', []))
        
        # Update bank stats
        if bank_code not in self.bank_stats:
            bank_info = get_bank_info(bank_code)
            self.bank_stats[bank_code] = {
                'name': bank_info['english_name'] if bank_info else bank_code,
                'newspapers_analyzed': 0,
                'newspapers_with_content': 0,
                'total_pages_found': 0
            }
        
        self.bank_stats[bank_code]['newspapers_analyzed'] += 1
        if result and result.get('bank_found'):
            self.bank_stats[bank_code]['newspapers_with_content'] += 1
            self.bank_stats[bank_code]['total_pages_found'] += len(result.get('relevant_pages', []))

    def get_summary_stats(self):
        """Get overall summary statistics"""
        total_content_found = sum(
            len([r for r in bank_results.values() if r['content_found']])
            for bank_results in self.results.values()
        )
        
        return {
            'total_combinations': self.total_combinations,
            'completed_combinations': self.completed_combinations,
            'combinations_with_content': total_content_found,
            'total_banks': len(set(self.results.keys())),
            'total_newspapers': len(set(self.newspaper_stats.keys())),
            'banks_analyzed': len(self.bank_stats),
            'newspapers_analyzed': len(self.newspaper_stats)
        }

    def get_bank_results(self, bank_code):
        """Get all results for a specific bank"""
        return self.results.get(bank_code, {})

    def get_newspaper_results(self, newspaper_script):
        """Get all bank results for a specific newspaper"""
        newspaper_results = {}
        for bank_code, bank_data in self.results.items():
            if newspaper_script in bank_data:
                newspaper_results[bank_code] = bank_data[newspaper_script]
        return newspaper_results

# ==================================================================================
# FIXED REPORT GENERATION FUNCTIONS
# ==================================================================================

def create_bank_analysis_table(doc, tracker, bank_code):
    """Create analysis table for a specific bank - FIXED VERSION"""
    bank_info = get_bank_info(bank_code)
    bank_name = bank_info['english_name'] if bank_info else bank_code
    
    doc.add_heading(f"{bank_name} - Analysis Summary", 2)
    
    # Create table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Headers
    headers = ['Newspaper', 'Status', 'Content Found', 'Pages Found', 'Error Details']
    header_cells = table.rows[0].cells
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add data rows - FIXED to show all newspapers
    bank_results = tracker.get_bank_results(bank_code)
    
    for script_name, display_name in NEWSPAPER_SCRIPTS.items():
        row_cells = table.add_row().cells
        
        if script_name in bank_results:
            data = bank_results[script_name]
            result = data['result']
            
            # Newspaper name
            row_cells[0].text = data['newspaper_name']
            
            # Status
            if data['success'] and result:
                row_cells[1].text = "✅ SUCCESS"
            else:
                row_cells[1].text = "❌ FAILED"
            
            # Content found - FIXED
            if result and result.get('bank_found'):
                row_cells[2].text = f"✅ Yes ({len(result.get('relevant_pages', []))} pages)"
            else:
                row_cells[2].text = "❌ No"
            
            # Pages found
            if result and result.get('relevant_pages'):
                pages_str = ', '.join([str(p+1) for p in result['relevant_pages']])
                row_cells[3].text = f"Pages: {pages_str}"
            else:
                row_cells[3].text = "None"
            
            # Error details
            if not data['success']:
                row_cells[4].text = "Analysis failed"
            elif not result.get('bank_found'):
                row_cells[4].text = "No content found"
            else:
                row_cells[4].text = "None"
        else:
            # Newspaper not attempted
            row_cells[0].text = display_name
            row_cells[1].text = "⚠️ NOT ATTEMPTED"
            row_cells[2].text = "❌ No"
            row_cells[3].text = "None"
            row_cells[4].text = "Not analyzed"
    
    doc.add_paragraph()


def resize_image_for_document(image_path, target_width=500, target_height=650):
    """Resize image to fit properly in Word document - ENHANCED VERSION"""
    try:
        with Image.open(image_path) as img:
            original_width, original_height = img.size
            
            # Conservative cropping - remove headers/footers
            max_height_for_page = 4700
            top_crop = 50
            
            if original_height > max_height_for_page:
                remaining_height_needed = max_height_for_page
                bottom_crop_position = top_crop + remaining_height_needed
                
                cropped_img = img.crop((0, top_crop, original_width, bottom_crop_position))
                resized_path = image_path.replace('.png', '_cropped.png')
                cropped_img.save(resized_path, 'PNG')
                
                return resized_path
            else:
                return image_path
    
    except Exception as e:
        print(f"   âŒ Error processing image: {e}")
        return image_path
    
def generate_analysis_charts(tracker, today):
    """Generate press releases and hits charts for the analysis"""
    print("Generating analysis charts...")
    
    # Prepare data for charts
    bank_names = []
    press_releases_count = []
    total_hits_count = []
    
    # Bank colors matching the charts - maintain consistent colors
    bank_colors = {
        'gulf_bank': '#8B0000',      # Dark red
        'nbk': '#000080',            # Navy blue  
        'cbk': '#228B22',            # Forest green
        'burgan_bank': '#FF8C00',    # Dark orange
        'boubyan_bank': '#DC143C',   # Crimson red
        'kfh': '#008B8B',            # Dark cyan
        'kib': '#2E8B57',            # Sea green
        'abk': '#DAA520',            # Goldenrod
        'abk': '#B8860B',            # Dark goldenrod (for AUB if different from ABK)
        'warba_bank': '#483D8B'      # Dark slate blue
    }
    
    # Process each bank from tracker
    for bank_code in tracker.bank_stats.keys():
        bank_info = get_bank_info(bank_code)
        bank_name = bank_info['english_name']
        
        # Shorten bank names for better chart display
        display_name = bank_name.replace('Bank', '').replace('of Kuwait', '').strip()
        if display_name == 'National': display_name = 'NBK'
        if display_name == 'Kuwait Finance House': display_name = 'KFH'
        if display_name == 'Commercial': display_name = 'CBK'
        if display_name == 'Kuwait International': display_name = 'KIB'
        if display_name == 'Al Ahli': display_name = 'ABK'
        
        bank_names.append(display_name)
        
        # Count press releases (newspapers with content for this bank)
        press_releases = tracker.bank_stats[bank_code].get('newspapers_with_content', 0)
        press_releases_count.append(press_releases)
        
        # Count total hits (total pages found across all newspapers)
        total_hits = tracker.bank_stats[bank_code].get('total_pages_found', 0)
        total_hits_count.append(total_hits)
    
    # Create charts directory
    charts_dir = "Charts"
    os.makedirs(charts_dir, exist_ok=True)
    
    # Set up matplotlib for better appearance
    plt.style.use('default')
    plt.rcParams['font.size'] = 10
    plt.rcParams['axes.labelsize'] = 12
    plt.rcParams['axes.titlesize'] = 14
    
    # Chart 1: Total Number of Press Releases
    fig1, ax1 = plt.subplots(figsize=(12, 6))
    
    # Get colors for each bank
    colors1 = [bank_colors.get(bank_code, '#808080') for bank_code in tracker.bank_stats.keys()]
    
    bars1 = ax1.bar(bank_names, press_releases_count, color=colors1, edgecolor='black', linewidth=0.5)
    ax1.set_title('Total Number of Press Releases', fontsize=16, fontweight='bold', pad=20)
    ax1.set_xlabel('Banks', fontsize=12, fontweight='bold')
    ax1.set_ylabel('Number of Press Releases', fontsize=12, fontweight='bold')
    ax1.grid(axis='y', alpha=0.3, linestyle='-', linewidth=0.5)
    ax1.set_ylim(0, max(press_releases_count + [1]) + 0.5)
    
    # Add value labels on bars
    for bar, value in zip(bars1, press_releases_count):
        if value > 0:
            ax1.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.02,
                    str(int(value)), ha='center', va='bottom', fontweight='bold')
    
    # Rotate x-axis labels for better readability
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    # Save chart 1
    chart1_path = os.path.join(charts_dir, f"Press_Releases_Chart_{today.strftime('%Y-%m-%d')}.png")
    plt.savefig(chart1_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    
    # Chart 2: Total Number of Hits
    fig2, ax2 = plt.subplots(figsize=(12, 6))
    
    # Get colors for each bank
    colors2 = [bank_colors.get(bank_code, '#808080') for bank_code in tracker.bank_stats.keys()]
    
    bars2 = ax2.bar(bank_names, total_hits_count, color=colors2, edgecolor='black', linewidth=0.5)
    ax2.set_title('Total Number of Hits', fontsize=16, fontweight='bold', pad=20)
    ax2.set_xlabel('Banks', fontsize=12, fontweight='bold')
    ax2.set_ylabel('Number of Hits', fontsize=12, fontweight='bold')
    ax2.grid(axis='y', alpha=0.3, linestyle='-', linewidth=0.5)
    ax2.set_ylim(0, max(total_hits_count + [1]) + 0.5)
    
    # Add value labels on bars
    for bar, value in zip(bars2, total_hits_count):
        if value > 0:
            ax2.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.1,
                    str(int(value)), ha='center', va='bottom', fontweight='bold')
    
    # Rotate x-axis labels for better readability
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    # Save chart 2
    chart2_path = os.path.join(charts_dir, f"Total_Hits_Chart_{today.strftime('%Y-%m-%d')}.png")
    plt.savefig(chart2_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    
    print(f"Charts generated:")
    print(f"  Press Releases: {chart1_path}")
    print(f"  Total Hits: {chart2_path}")
    
    return chart1_path, chart2_path

# Add this new function to main.py (after the generate_analysis_charts function)

def create_analytics_report(tracker, today):
    """Create separate analytics report with charts and statistics only"""
    print("Creating separate analytics report...")
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Title
    title = doc.add_heading(f"Banking News Analytics Report ({today.strftime('%Y-%m-%d')})", 1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Executive Summary
    stats = tracker.get_summary_stats()
    total_content_instances = sum(
        bank_data.get('newspapers_with_content', 0) 
        for bank_data in tracker.bank_stats.values()
    )
    
    subtitle_text = f"Statistical Analysis of {stats['total_banks']} banks across {stats['total_newspapers']} newspapers"
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    
    # Generate charts
    try:
        chart1_path, chart2_path = generate_analysis_charts(tracker, today)
        
        # Charts Section
        doc.add_heading("Visual Analytics", 2)
        
        # Press Releases Chart
        if chart1_path and os.path.exists(chart1_path):
            doc.add_heading("Total Number of Press Releases by Bank", 3)
            chart_para = doc.add_paragraph()
            chart_run = chart_para.add_run()
            try:
                chart_run.add_picture(chart1_path, width=Inches(6.5))
                chart_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add description
                desc_para = doc.add_paragraph()
                desc_para.add_run("This chart shows the number of newspapers that contained press releases or news coverage for each bank. A higher value indicates broader media coverage across multiple publications.").italic = True
                desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
            except Exception as e:
                print(f"Error adding press releases chart: {e}")
                doc.add_paragraph("Press releases chart could not be generated.")
        
        doc.add_page_break()
        
        # Total Hits Chart
        if chart2_path and os.path.exists(chart2_path):
            doc.add_heading("Total Number of Content Hits by Bank", 3)
            chart_para = doc.add_paragraph()
            chart_run = chart_para.add_run()
            try:
                chart_run.add_picture(chart2_path, width=Inches(6.5))
                chart_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Add description
                desc_para = doc.add_paragraph()
                desc_para.add_run("This chart displays the total number of individual content instances (pages) found for each bank across all analyzed newspapers. This includes multiple mentions within the same publication.").italic = True
                desc_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
            except Exception as e:
                print(f"Error adding hits chart: {e}")
                doc.add_paragraph("Content hits chart could not be generated.")
    
    except Exception as e:
        print(f"Error generating charts: {e}")
        doc.add_paragraph("Charts could not be generated due to an error.")
    
    doc.add_page_break()
    
    # Detailed Statistics Section
    doc.add_heading("Detailed Statistics", 2)
    
    # Overall Analysis Summary
    doc.add_heading("Analysis Overview", 3)
    overview_stats = [
        f"Total bank-newspaper combinations analyzed: {stats['total_combinations']}",
        f"Successfully completed analyses: {stats['completed_combinations']}",
        f"Combinations with content found: {stats['combinations_with_content']}",
        f"Overall success rate: {(stats['combinations_with_content']/stats['total_combinations']*100):.1f}%" if stats['total_combinations'] > 0 else "Overall success rate: 0%",
        f"Banks analyzed: {stats['total_banks']}",
        f"Newspapers analyzed: {stats['total_newspapers']}"
    ]
    
    for stat in overview_stats:
        doc.add_paragraph(stat, style='List Bullet')
    
    # Bank Performance Table
    doc.add_heading("Bank Performance Summary", 3)
    
    bank_table = doc.add_table(rows=1, cols=5)
    bank_table.style = 'Table Grid'
    
    # Headers
    headers = ['Bank', 'Press Releases', 'Total Content Hits', 'Coverage Rate', 'Performance Score']
    header_cells = bank_table.rows[0].cells
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add bank data rows
    total_newspapers = len(tracker.newspaper_stats)
    
    for bank_code in tracker.bank_stats.keys():
        bank_stats = tracker.bank_stats.get(bank_code, {})
        bank_name = bank_stats.get('name', bank_code)
        newspapers_with_content = bank_stats.get('newspapers_with_content', 0)
        total_pages_found = bank_stats.get('total_pages_found', 0)
        coverage_rate = (newspapers_with_content / total_newspapers) * 100 if total_newspapers > 0 else 0
        
        # Calculate performance score (weighted combination of coverage rate and content volume)
        performance_score = (coverage_rate * 0.7) + (min(total_pages_found, 20) * 1.5)  # Cap content hits at 20 for scoring
        
        row_cells = bank_table.add_row().cells
        row_cells[0].text = bank_name
        row_cells[1].text = str(newspapers_with_content)
        row_cells[2].text = str(total_pages_found)
        row_cells[3].text = f"{coverage_rate:.1f}%"
        row_cells[4].text = f"{performance_score:.1f}"
        
        # Color coding for performance score
        if performance_score > 75:
            row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif performance_score > 40:
            row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            row_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    doc.add_paragraph()
    
    # Newspaper Performance Table
    doc.add_heading("Newspaper Coverage Analysis", 3)
    
    newspaper_table = doc.add_table(rows=1, cols=4)
    newspaper_table.style = 'Table Grid'
    
    # Headers
    headers = ['Newspaper', 'Banks Covered', 'Total Content Found', 'Coverage Effectiveness']
    header_cells = newspaper_table.rows[0].cells
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add newspaper data rows
    total_banks = len(tracker.bank_stats)
    
    for script_name, newspaper_data in tracker.newspaper_stats.items():
        newspaper_name = newspaper_data['name']
        banks_with_content = newspaper_data['banks_with_content']
        total_pages_found = newspaper_data['total_pages_found']
        effectiveness = (banks_with_content / total_banks) * 100 if total_banks > 0 else 0
        
        row_cells = newspaper_table.add_row().cells
        row_cells[0].text = newspaper_name
        row_cells[1].text = f"{banks_with_content}/{total_banks}"
        row_cells[2].text = str(total_pages_found)
        row_cells[3].text = f"{effectiveness:.1f}%"
        
        # Color coding for effectiveness
        if effectiveness > 60:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif effectiveness > 30:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    doc.add_paragraph()
    
    # Key Insights Section
    doc.add_heading("Key Insights", 3)
    
    # Calculate insights
    max_coverage_bank = max(tracker.bank_stats.items(), key=lambda x: x[1].get('newspapers_with_content', 0))
    max_content_bank = max(tracker.bank_stats.items(), key=lambda x: x[1].get('total_pages_found', 0))
    best_newspaper = max(tracker.newspaper_stats.items(), key=lambda x: x[1].get('banks_with_content', 0))
    
    insights = [
        f"Most covered bank: {max_coverage_bank[1].get('name', 'Unknown')} with coverage in {max_coverage_bank[1].get('newspapers_with_content', 0)} newspapers",
        f"Bank with most content: {max_content_bank[1].get('name', 'Unknown')} with {max_content_bank[1].get('total_pages_found', 0)} content instances",
        f"Most comprehensive newspaper: {best_newspaper[1].get('name', 'Unknown')} covering {best_newspaper[1].get('banks_with_content', 0)} banks",
        f"Average content per bank: {sum(bank_data.get('total_pages_found', 0) for bank_data in tracker.bank_stats.values()) / len(tracker.bank_stats):.1f} instances" if tracker.bank_stats else "No data available"
    ]
    
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')
    
    # Methodology Notes
    doc.add_heading("Methodology", 3)
    methodology_notes = [
        "Press Releases: Count of newspapers containing content for each bank",
        "Content Hits: Total number of individual content instances (pages/articles) found",
        "Coverage Rate: Percentage of analyzed newspapers that contained bank content",
        "Performance Score: Weighted combination of coverage breadth and content volume",
        "Analysis conducted using Vision API with LLM validation for accuracy"
    ]
    
    for note in methodology_notes:
        doc.add_paragraph(note, style='List Bullet')
    
    # Footer
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    footer_para = doc.add_paragraph("Analytics report generated on: ")
    footer_para.add_run(f"{datetime.now(KW).strftime('%Y-%m-%d %H:%M')}").italic = True
    
    # Save report
    reports_dir = "Reports"
    os.makedirs(reports_dir, exist_ok=True)
    report_path = os.path.join(reports_dir, f"Banking_Analytics_Report_{today.strftime('%Y-%m-%d')}.docx")
    doc.save(report_path)
    
    print(f"Analytics report saved: {report_path}")
    return report_path

def create_individual_bank_report(bank_code, tracker, today):
    """Create detailed individual bank report with comprehensive English content"""
    bank_info = get_bank_info(bank_code)
    bank_name = bank_info['english_name'] if bank_info else bank_code
    
    print(f"Creating detailed individual report for {bank_name}...")
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Title
    title = doc.add_heading(f"{bank_name} Detailed News Analysis ({today.strftime('%Y-%m-%d')})", 1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Subtitle
    bank_results = tracker.get_bank_results(bank_code)
    newspapers_with_content = len([r for r in bank_results.values() if r['content_found']])
    
    subtitle_text = f"Comprehensive analysis found {bank_name} content in {newspapers_with_content} out of {len(bank_results)} newspapers analyzed"
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    
    # Analysis table for this bank
    create_bank_analysis_table(doc, tracker, bank_code)
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    
    # Detailed content sections with enhanced information
    content_found = False
    newspaper_count = 0
    
    # Get newspapers with content
    newspapers_with_content_list = []
    for script_name, result_data in bank_results.items():
        result = result_data['result']
        if result and result.get('bank_found'):
            newspapers_with_content_list.append((script_name, result_data))
    
    for script_name, result_data in newspapers_with_content_list:
        result = result_data['result']
        content_found = True
        newspaper_count += 1
        
        # Add page break before each newspaper (except the first one)
        if newspaper_count > 1:
            doc.add_page_break()
        
        # Newspaper heading
        doc.add_heading(f"{result_data['newspaper_name']} - {bank_name} Coverage", 2)
        
        # Content Type and Confidence
        doc.add_heading("Content Analysis", 3)
        analysis_para = doc.add_paragraph()
        analysis_para.add_run("Content Type: ").bold = True
        analysis_para.add_run(f"{result.get('content_type', 'Not specified')}\n")
        analysis_para.add_run("Detection Method: ").bold = True
        analysis_para.add_run(f"{result.get('detection_method', 'Vision API')}\n")
        if result.get('confidence'):
            analysis_para.add_run("Confidence Level: ").bold = True
            analysis_para.add_run(f"{result.get('confidence', 0.0):.1%}")
        
        # Pages with content
        if result.get('relevant_pages'):
            page_nums = [str(p+1) for p in sorted(result['relevant_pages'])]
            doc.add_heading("Pages with Content", 3)
            doc.add_paragraph(f"Content found on page(s): {', '.join(page_nums)}")
        
        # Detailed Summary
        if result.get('summary') and result['summary'] != "No summary available":
            doc.add_heading("Detailed Summary", 3)
            doc.add_paragraph(result['summary'])
        
        # Headlines and Key Points
        if result.get('highlights'):
            doc.add_heading("Headlines and Key Points", 3)
            for highlight in result['highlights']:
                clean_highlight = highlight.lstrip('•-* ').strip()
                # Remove any existing numbering
                clean_highlight = re.sub(r'^\d+\.\s*', '', clean_highlight)
                doc.add_paragraph(clean_highlight, style='List Number')
            # for i, highlight in enumerate(result['highlights'], 1):
            #     clean_highlight = highlight.lstrip('•-* ').strip()
            #     #clean_highlight = re.sub(r'^\d+\.\s*', '', clean_highlight)
            #     doc.add_paragraph(f"{i}. {clean_highlight}", style='List Number')
        
        # Key Details (if available from Vision API)
        if result.get('key_details'):
            doc.add_heading("Key Details", 3)
            for detail in result['key_details']:
                doc.add_paragraph(detail, style='List Bullet')
        
        # Content Context (if available)
        if result.get('context_snippets'):
            doc.add_heading("Content Context", 3)
            for snippet in result['context_snippets']:
                context_para = doc.add_paragraph()
                context_para.add_run('"').italic = True
                context_para.add_run(snippet[:200] + "..." if len(snippet) > 200 else snippet).italic = True
                context_para.add_run('"').italic = True
        
        # Screenshots section (unchanged)
        if result.get('page_screenshots') and result.get('relevant_pages'):
            doc.add_heading("Page Screenshots", 3)
            
            print(f"   Checking screenshots for {result_data['newspaper_name']}...")
            
            screenshot_count = 0
            total_screenshots = len(result['relevant_pages'])
            screenshots_added = 0
            
            for page_num in sorted(result['relevant_pages']):
                screenshot_path = result['page_screenshots'].get(page_num)
                screenshot_count += 1
                
                if screenshot_path and os.path.exists(screenshot_path):
                    print(f"   Adding screenshot for page {page_num+1}: {screenshot_path}")
                    
                    # Add page label
                    page_para = doc.add_paragraph()
                    page_run = page_para.add_run(f"Page {page_num+1}")
                    page_run.bold = True
                    page_run.font.size = Pt(14)
                    page_para.paragraph_format.keep_with_next = True
                    
                    # Add the screenshot
                    try:
                        processed_path = resize_image_for_document(screenshot_path)
                        image_para = doc.add_paragraph()
                        image_run = image_para.add_run()
                        image_run.add_picture(processed_path, width=Inches(4.2))
                        image_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        screenshots_added += 1
                        
                        if screenshot_count < total_screenshots:
                            doc.add_page_break()
                        elif newspaper_count < len(newspapers_with_content_list):
                            pass
                        else:
                            spacing_para = doc.add_paragraph()
                            spacing_para.space_after = Pt(24)
                        
                    except Exception as e:
                        print(f"   Error adding image for page {page_num+1}: {e}")
                        error_para = doc.add_paragraph(f"[Image could not be added: {os.path.basename(screenshot_path)}]")
                        error_para.runs[0].italic = True
    
    # No content message if needed
    if not content_found:
        doc.add_heading("No Content Found", 2)
        doc.add_paragraph(f"No {bank_name} content was found in any of the analyzed newspapers today.")
    
    # Footer
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    footer_para = doc.add_paragraph("Detailed report generated on: ")
    footer_para.add_run(f"{datetime.now(KW).strftime('%Y-%m-%d %H:%M')}").italic = True
    
    # Save report
    reports_dir = "Reports"
    os.makedirs(reports_dir, exist_ok=True)
    report_path = os.path.join(reports_dir, f"{bank_name.replace(' ', '_')}_Detailed_Analysis_{today.strftime('%Y-%m-%d')}.docx")
    doc.save(report_path)
    
    print(f"{bank_name} detailed report saved: {report_path}")
    return report_path



def create_headlines_report(tracker, today):
    """Create consolidated headlines report with headlines only (no summaries) - NO CHARTS"""
    print("Creating consolidated headlines report...")
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Title
    title = doc.add_heading(f"Banking News Headlines Report ({today.strftime('%Y-%m-%d')})", 1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Executive Summary
    stats = tracker.get_summary_stats()
    total_content_instances = sum(
        bank_data.get('newspapers_with_content', 0) 
        for bank_data in tracker.bank_stats.values()
    )
    
    subtitle_text = f"Analysis of {stats['total_banks']} banks across {stats['total_newspapers']} newspapers • {total_content_instances} content instances found"
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    
    # Enhanced statistics WITHOUT charts (original version)
    create_enhanced_statistics(doc, tracker)
    
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    
    # Headlines by bank - HEADLINES ONLY, NO SUMMARIES
    doc.add_heading("Headlines by Bank", 2)
    
    for bank_code in tracker.bank_stats.keys():
        bank_info = get_bank_info(bank_code)
        bank_name = bank_info['english_name'] if bank_info else bank_code
        
        doc.add_heading(bank_name, 3)
        
        bank_results = tracker.get_bank_results(bank_code)
        bank_has_content = False
        
        for script_name, result_data in bank_results.items():
            result = result_data['result']
            if result and result.get('bank_found') and result.get('highlights'):
                bank_has_content = True
                
                # Newspaper sub-heading
                newspaper_para = doc.add_paragraph()
                newspaper_run = newspaper_para.add_run(f"{result_data['newspaper_name']}:")
                newspaper_run.bold = True
                newspaper_run.font.size = Pt(12)
                
                # Add headlines ONLY (no summaries)
                for highlight in result['highlights']:
                    clean_highlight = highlight.lstrip('•-* ').strip()
                    doc.add_paragraph(clean_highlight, style='List Bullet')
        
        if not bank_has_content:
            doc.add_paragraph(f"No {bank_name} news found today.")
        
        doc.add_paragraph()  # Add spacing between banks
    
    # Analysis Summary
    doc.add_heading("Analysis Summary", 2)
    
    summary_points = [
        f"Analyzed {stats['total_combinations']} bank-newspaper combinations",
        f"Successfully completed {stats['completed_combinations']} analyses",
        f"Found banking content in {stats['combinations_with_content']} combinations",
        f"Generated detailed reports for all {stats['total_banks']} banks"
    ]
    
    for point in summary_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Footer
    doc.add_paragraph().add_run("_" * 80).font.size = Pt(12)
    footer_para = doc.add_paragraph("Report generated on: ")
    footer_para.add_run(f"{datetime.now(KW).strftime('%Y-%m-%d %H:%M')}").italic = True
    
    # Save report
    reports_dir = "Reports"
    os.makedirs(reports_dir, exist_ok=True)
    report_path = os.path.join(reports_dir, f"Banking_News_Headlines_Report_{today.strftime('%Y-%m-%d')}.docx")
    doc.save(report_path)
    
    print(f"Headlines report saved: {report_path}")
    return report_path

def calculate_content_quality_score(result, raw_text_context=""):
    """Simplified quality scoring - less restrictive"""
    score = 50.0  # Start with base score of 50
    
    # Check if we have basic content
    summary = result.get('summary', '')
    headlines = result.get('highlights', [])
    
    if summary and len(summary) > 20:
        score += 20
    
    if headlines and len(headlines) > 0:
        score += 20
    
    # Bonus for substantive content
    if any(word in summary.lower() for word in ['announces', 'reports', 'launches', 'achieves', 'posts']):
        score += 10
    
    return min(score, 100.0)
    

async def llm_validate_bank_content(bank_code, result, raw_text_chunk):
    """Balanced LLM validation to reduce false positives while keeping genuine news"""
    
    bank_info = get_bank_info(bank_code)
    
    validation_prompt = f"""
You are validating banking news content for {bank_info['english_name']}.

CONTENT TO VALIDATE:
Summary: {result.get('summary', '')}
Headlines: {result.get('highlights', [])}
Raw Text Sample: {raw_text_chunk[:2000]}

VALIDATION CRITERIA:
1. Is there genuine news content about {bank_info['english_name']}?
2. Does the bank have a meaningful role in the story (not just a passing mention)?
3. Is this banking/financial news rather than just legal notices or generic market stats?

REJECT ONLY IF:
- Bank is only mentioned in passing statistics without any specific action
- Content is purely legal/regulatory notice with no news value
- Bank is mentioned only as "among banks" or "facilitated by" without being subject
- Content is advertisement or sponsorship mention only

ACCEPT IF:
- Bank announces, launches, reports, or does something specific
- Bank is featured in financial results or business news
- Bank participates in events, partnerships, or initiatives
- Bank has executive changes or strategic decisions

Respond with JSON only:
{{
  "is_valid": true/false,
  "confidence": 0.0-1.0,
  "reason": "Brief explanation"
}}
"""

    try:
        response = await asyncio.to_thread(
            openai_client.chat.completions.create,
            model="gpt-4o-mini",  # Back to the original model
            messages=[{"role": "user", "content": validation_prompt}],
            temperature=0.1,
            max_tokens=200
        )
        
        result_text = response.choices[0].message.content.strip()
        
        # Extract JSON from response
        json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
        if json_match:
            validation_result = json.loads(json_match.group())
        else:
            validation_result = json.loads(result_text)
        
        return (
            validation_result.get('is_valid', False),
            validation_result.get('confidence', 0.0),
            validation_result.get('reason', 'No reason provided')
        )
        
    except Exception as e:
        print(f"LLM validation error: {e}")
        return True, 1.0, f"Validation failed: {e}"  # Default to accepting when validation fails
    
def create_enhanced_statistics(doc, tracker):
    """Create enhanced statistics section with detailed breakdowns (NO CHARTS)"""
    doc.add_heading("Analysis Statistics", 2)
    
    # 1. Bank Coverage Summary Table
    doc.add_heading("Bank Coverage Summary", 3)
    
    bank_table = doc.add_table(rows=1, cols=4)
    bank_table.style = 'Table Grid'
    
    # Headers
    headers = ['Bank', 'Newspapers with Content', 'Total Pages Found', 'Coverage Rate']
    header_cells = bank_table.rows[0].cells
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        header_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add bank data rows
    total_newspapers = len(tracker.newspaper_stats)
    
    for bank_code in tracker.bank_stats.keys():
        bank_stats = tracker.bank_stats.get(bank_code, {})
        bank_name = bank_stats.get('name', bank_code)
        newspapers_with_content = bank_stats.get('newspapers_with_content', 0)
        total_pages_found = bank_stats.get('total_pages_found', 0)
        coverage_rate = (newspapers_with_content / total_newspapers) * 100 if total_newspapers > 0 else 0
        
        row_cells = bank_table.add_row().cells
        row_cells[0].text = bank_name
        row_cells[1].text = f"{newspapers_with_content}/{total_newspapers}"
        row_cells[2].text = str(total_pages_found)
        row_cells[3].text = f"{coverage_rate:.1f}%"
        
        # Color coding for coverage rate
        if coverage_rate > 50:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif coverage_rate > 25:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            row_cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    doc.add_paragraph()
    

# ==================================================================================
# FIXED NEWSPAPER SCRIPT EXECUTION WITH ENHANCED SUBPROCESS APPROACH
# ==================================================================================

def run_newspaper_analysis_subprocess(newspaper_id: str, selected_banks: List[str]) -> Dict:
    """FIXED: Run newspaper analysis using subprocess - NOT ASYNC"""
    try:
        print(f"🔧 Running {newspaper_id} analysis via subprocess to avoid loop conflicts...")
        print(f"📑 Selected banks for this analysis: {selected_banks}")
        
        # Check if script exists
        script_path = f"{newspaper_id}.py"
        if not os.path.exists(script_path):
            print(f"❌ Script not found: {script_path}")
            return {bank_code: None for bank_code in selected_banks}
        
        # Create temporary directory for runner scripts
        temp_dir = tempfile.mkdtemp()
        runner_path = os.path.join(temp_dir, f"runner_{newspaper_id}_{uuid.uuid4().hex[:8]}.py")
        
        # FIXED: Create a subprocess runner script that PRESERVES SCREENSHOTS
        runner_code = f"""
import asyncio
import sys
import json
import os
import fitz

# Set event loop policy to avoid uvloop conflicts
asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())

# Add current directory to path
sys.path.insert(0, '{os.getcwd()}')

async def run_analysis():
    try:
        import {newspaper_id}
        
        if hasattr({newspaper_id}, 'analyze_all_banks'):
            results = await {newspaper_id}.analyze_all_banks()
            
            # NEW: Extract raw text for validation
            pdf_path = None
            for result in results.values():
                if result and result.get('pdf_path'):
                    pdf_path = result['pdf_path']
                    break
            
            # Extract raw text from relevant pages
            page_texts = {{}}
            if pdf_path and os.path.exists(pdf_path):
                try:
                    doc = fitz.open(pdf_path)
                    for result in results.values():
                        if result and result.get('relevant_pages'):
                            for page_num in result['relevant_pages']:
                                if page_num not in page_texts:
                                    page = doc.load_page(page_num)
                                    page_texts[page_num] = page.get_text()
                    doc.close()
                except Exception as e:
                    print(f"Error extracting page text: {{e}}")
            
            # Add page texts to results
            for bank_code, result in results.items():
                if result and result.get('relevant_pages'):
                    result['page_texts'] = {{}}
                    for page_num in result['relevant_pages']:
                        result['page_texts'][page_num] = page_texts.get(page_num, '')
            
            # Filter and return results...
            filtered_results = {{}}
            selected_banks = {selected_banks}
            for bank_code in selected_banks:
                if bank_code in results:
                    filtered_results[bank_code] = results[bank_code]
            
            return filtered_results
        else:
            return {{}}
            
    except Exception as e:
        print(f"Error in subprocess analysis: {{e}}")
        return {{}}
    
if __name__ == "__main__":
    try:
        # Run with fresh event loop
        results = asyncio.run(run_analysis())
        print("RESULTS_START")
        print(json.dumps(results, default=str))
        print("RESULTS_END")
    except Exception as e:
        print(f"Subprocess error: {{e}}")
        print("RESULTS_START")
        print(json.dumps({{}}, default=str))
        print("RESULTS_END")
"""
        
        # Write the runner script to a temporary file
        with open(runner_path, 'w', encoding='utf-8') as f:
            f.write(runner_code)
        
        try:
            # Run the subprocess with timeout
            print(f"🚀 Starting subprocess for {newspaper_id}...")
            
            result = subprocess.run([
                sys.executable, runner_path
            ], 
            capture_output=True, 
            text=True, 
            timeout=1800,  # 30 minute timeout
            env={**os.environ, 'PYTHONPATH': os.getcwd()}
            )
            
            print(f"📊 Subprocess completed for {newspaper_id}")
            print(f"   Return code: {result.returncode}")
            print(f"   Stdout length: {len(result.stdout)} chars")
            print(f"   Stderr length: {len(result.stderr)} chars")
            
            if result.stderr:
                print(f"   Stderr preview: {result.stderr[:200]}...")
            
            # Clean up runner script and directory
            try:
                os.remove(runner_path)
                os.rmdir(temp_dir)
            except Exception as cleanup_error:
                print(f"⚠️ Cleanup warning: {cleanup_error}")
            
            if result.returncode == 0:
                # Parse results from stdout
                output = result.stdout
                if "RESULTS_START" in output and "RESULTS_END" in output:
                    start_idx = output.find("RESULTS_START") + len("RESULTS_START")
                    end_idx = output.find("RESULTS_END")
                    results_json = output[start_idx:end_idx].strip()
                    
                    try:
                        results = json.loads(results_json)
                        print(f"✅ Successfully parsed results for {newspaper_id}")
                        print(f"   Results keys: {list(results.keys())}")
                        
                        # FIXED: Validate results structure AND screenshot paths
                        valid_results = {}
                        for bank_code in selected_banks:
                            if bank_code in results and results[bank_code]:
                                result_data = results[bank_code]
                                
                                # Check if screenshots exist and log details
                                if result_data.get('page_screenshots'):
                                    screenshot_count = 0
                                    valid_screenshots = {}
                                    for page_num, screenshot_path in result_data['page_screenshots'].items():
                                        if os.path.exists(screenshot_path):
                                            valid_screenshots[int(page_num)] = screenshot_path
                                            screenshot_count += 1
                                            print(f"   📷 Valid screenshot: {screenshot_path}")
                                        else:
                                            print(f"   ⚠️ Screenshot missing: {screenshot_path}")
                                    
                                    result_data['page_screenshots'] = valid_screenshots
                                    print(f"   📷 {bank_code}: {screenshot_count} valid screenshots found")
                                
                                # Log content detection
                                bank_found = result_data.get('bank_found', False)
                                relevant_pages = result_data.get('relevant_pages', [])
                                print(f"   🔍 {bank_code}: bank_found={bank_found}, pages={len(relevant_pages)}")

                                valid_results[bank_code] = result_data
                                print(f"   ✅ {bank_code}: Valid result found")
                            else:
                                valid_results[bank_code] = None
                                print(f"   ❌ {bank_code}: No valid result")
                        
                        return valid_results
                        
                    except json.JSONDecodeError as e:
                        print(f"❌ Failed to parse JSON results: {e}")
                        print(f"Raw results preview: {results_json[:300]}...")
                        return {bank_code: None for bank_code in selected_banks}
                else:
                    print(f"❌ No results markers found in output")
                    print(f"Output preview: {output[:500]}...")
                    return {bank_code: None for bank_code in selected_banks}
            else:
                print(f"❌ Subprocess failed with return code {result.returncode}")
                if result.stderr:
                    print(f"Error details: {result.stderr}")
                return {bank_code: None for bank_code in selected_banks}
                
        except subprocess.TimeoutExpired:
            print(f"⏰ Subprocess timeout for {newspaper_id}")
            return {bank_code: None for bank_code in selected_banks}
        except Exception as e:
            print(f"❌ Subprocess error for {newspaper_id}: {e}")
            return {bank_code: None for bank_code in selected_banks}
            
    except Exception as e:
        print(f"❌ Error in run_newspaper_analysis_subprocess for {newspaper_id}: {e}")
        return {bank_code: None for bank_code in selected_banks}

# FIXED: Enhanced results formatter
def format_results_for_frontend(tracker, selected_banks, selected_newspapers):
    """FIXED: Convert tracker results to frontend-compatible format - following master orchestrator pattern"""
    frontend_results = {}
    
    try:
        for bank_code in selected_banks:
            bank_results = tracker.get_bank_results(bank_code)
            bank_info = get_bank_info(bank_code)
            
            for newspaper_id in selected_newspapers:
                if newspaper_id in bank_results:
                    result_data = bank_results[newspaper_id]
                    result = result_data.get('result')
                    result_key = f"{bank_code}_{newspaper_id}"
                    
                    if result:
                        # FIXED: Handle date properly (like master orchestrator)
                        try:
                            if hasattr(result.get('date'), 'isoformat'):
                                date_string = result['date'].isoformat()
                            else:
                                date_string = str(result.get('date', datetime.now(KW).isoformat()))
                        except:
                            date_string = datetime.now(KW).isoformat()
                        
                        frontend_results[result_key] = {
                            "bank": bank_info['english_name'],
                            "newspaper": result_data['newspaper_name'],
                            "hasContent": result.get('bank_found', False),
                            "pages": len(result.get('relevant_pages', [])),
                            "highlights": result.get('highlights', [])[:5],
                            "error": False,
                            "summary": result.get('summary', ''),
                            "relevant_pages": result.get('relevant_pages', []),
                            "analysis_date": date_string
                        }
                    else:
                        frontend_results[result_key] = {
                            "bank": bank_info['english_name'],
                            "newspaper": result_data['newspaper_name'],
                            "hasContent": False,
                            "pages": 0,
                            "highlights": [],
                            "error": True,
                            "errorMessage": f"Failed to analyze {result_data['newspaper_name']} for {bank_info['english_name']}",
                            "summary": "",
                            "relevant_pages": [],
                            "analysis_date": datetime.now(KW).isoformat()
                        }
                else:
                    # Newspaper not analyzed for this bank
                    result_key = f"{bank_code}_{newspaper_id}"
                    newspaper_name = NEWSPAPER_SCRIPTS.get(newspaper_id, newspaper_id)
                    
                    frontend_results[result_key] = {
                        "bank": bank_info['english_name'],
                        "newspaper": newspaper_name,
                        "hasContent": False,
                        "pages": 0,
                        "highlights": [],
                        "error": True,
                        "errorMessage": f"{newspaper_name} not analyzed for {bank_info['english_name']}",
                        "summary": "",
                        "relevant_pages": [],
                        "analysis_date": datetime.now(KW).isoformat()
                    }
        
        return frontend_results
        
    except Exception as e:
        print(f"❌ Error formatting results for frontend: {e}")
        import traceback
        traceback.print_exc()
        return {}

async def load_and_run_newspaper_script(script_name, display_name, tracker, selected_banks, task_id):
    """ENHANCED: Load and run a newspaper analysis script with LLM validation"""
    print(f"\n📰 Starting analysis for {display_name}...")
    print(f"📋 Analyzing for selected banks: {[get_bank_info(bank)['english_name'] for bank in selected_banks]}")
    
    # Send immediate update when starting newspaper
    if task_id in active_tasks:
        active_tasks[task_id]["progress"].update({
            "current_task": f"📰 Starting {display_name} analysis...",
            "phase": "downloading"
        })
        print(f"📡 WebSocket update: Starting {display_name}")
    
    try:
        # Check if script exists
        script_path = f"{script_name}.py"
        if not os.path.exists(script_path):
            error_msg = f"Could not find script: {script_name}.py"
            print(f"❌ {error_msg}")
            
            # Add failed results for all banks
            for bank_code in selected_banks:
                tracker.add_result(bank_code, script_name, display_name, None)
            return
        
        # Send update when downloading starts
        if task_id in active_tasks:
            active_tasks[task_id]["progress"].update({
                "current_task": f"📥 Downloading {display_name}...",
                "phase": "downloading"
            })
            print(f"📡 WebSocket update: Downloading {display_name}")
            await asyncio.sleep(0.5)
        
        # Send update when analysis starts
        if task_id in active_tasks:
            active_tasks[task_id]["progress"].update({
                "current_task": f"🔍 Analyzing {display_name} for {len(selected_banks)} banks...",
                "phase": "analyzing"
            })
            print(f"📡 WebSocket update: Analyzing {display_name}")
            await asyncio.sleep(0.5)
        
        # Run the newspaper analysis
        print(f"   🧠 Using subprocess approach for {display_name}")
        results = run_newspaper_analysis_subprocess(script_name, selected_banks)
        
        # Send update when processing results
        if task_id in active_tasks:
            active_tasks[task_id]["progress"].update({
                "current_task": f"📊 Processing {display_name} results...",
                "phase": "analyzing"
            })
            print(f"📡 WebSocket update: Processing {display_name} results")
            await asyncio.sleep(0.5)
        
        # Add results for each selected bank WITH LLM VALIDATION
        content_found_count = 0
        llm_validation_enabled = VALIDATION_CONFIG.get('llm_validation_enabled', True)
        
        for i, bank_code in enumerate(selected_banks):
            if bank_code in results and results[bank_code]:
                result = results[bank_code]
                result['newspaper_name'] = display_name
                
                bank_name = get_bank_info(bank_code)['english_name']
                
                # Enhanced LLM-based validation with quality scoring
                # Balanced LLM-based validation
                if llm_validation_enabled and result.get('bank_found', False):
                    print(f"   🤖 Running balanced LLM validation for {bank_name}...")
                    
                    # Get raw text for validation
                    raw_text = ""
                    if result.get('page_texts'):
                        raw_text = " ".join(result['page_texts'].values())
                    elif result.get('analysis_text'):
                        raw_text = result['analysis_text']
                    else:
                        raw_text = f"{result.get('summary', '')} {' '.join(result.get('highlights', []))}"
                    
                    # Only apply quality filter for very poor content
                    quality_score = calculate_content_quality_score(result, raw_text)
                    
                    if quality_score < 30.0:  # Much lower threshold
                        print(f"   ❌ QUALITY FILTER FAILED: {bank_name} (Score: {quality_score:.1f}/100)")
                        filtered_result = result.copy()
                        filtered_result['bank_found'] = False
                        filtered_result['relevant_pages'] = []
                        filtered_result['analysis_text'] = f"Content filtered: very low quality score {quality_score:.1f}/100"
                        filtered_result['summary'] = "Content filtered as very low quality"
                        filtered_result['highlights'] = []
                        filtered_result['page_screenshots'] = {}
                        filtered_result['validation_passed'] = False
                        filtered_result['validation_reason'] = f"Quality too low: {quality_score:.1f}/100"
                        filtered_result['validation_method'] = 'quality_filter'
                        
                        tracker.add_result(bank_code, script_name, display_name, filtered_result)
                        status = "❌ FILTERED BY QUALITY"
                    else:
                        # Apply balanced LLM validation
                        is_valid, confidence_score, validation_reason = await llm_validate_bank_content(
                            bank_code, result, raw_text
                        )
                        
                        # Accept if confidence is reasonable OR if validation failed (give benefit of doubt)
                        if is_valid and confidence_score >= 0.3:  # Lower threshold
                            print(f"   ✅ VALIDATION PASSED: {bank_name} (Confidence: {confidence_score:.2f}, Quality: {quality_score:.1f})")
                            
                            result['validation_score'] = confidence_score
                            result['validation_passed'] = True
                            result['validation_reason'] = f"Balanced validation passed (Quality: {quality_score:.1f}/100)"
                            result['validation_method'] = 'balanced_llm'
                            result['quality_score'] = quality_score
                            
                            tracker.add_result(bank_code, script_name, display_name, result)
                            content_found_count += 1
                            status = "✅ VALIDATED"
                            
                            print(f"   📊 Content details: {len(result.get('relevant_pages', []))} pages, {len(result.get('highlights', []))} highlights")
                        else:
                            print(f"   ❌ VALIDATION FAILED: {bank_name} (Confidence: {confidence_score:.2f})")
                            print(f"      Reason: {validation_reason}")
                            
                            filtered_result = result.copy()
                            filtered_result['bank_found'] = False
                            filtered_result['relevant_pages'] = []
                            filtered_result['analysis_text'] = f"Content filtered by LLM: {validation_reason}"
                            filtered_result['summary'] = "Content filtered as false positive"
                            filtered_result['highlights'] = []
                            filtered_result['page_screenshots'] = {}
                            filtered_result['validation_score'] = confidence_score
                            filtered_result['validation_passed'] = False
                            filtered_result['validation_reason'] = validation_reason
                            filtered_result['validation_method'] = 'balanced_llm'
                            filtered_result['quality_score'] = quality_score
                            
                            tracker.add_result(bank_code, script_name, display_name, filtered_result)
                            status = "❌ FILTERED BY LLM"
                
                else:
                    # No LLM validation or no content found
                    bank_found = result.get('bank_found', False)
                    
                    if not llm_validation_enabled:
                        # Validation disabled - add as before
                        result['validation_score'] = 1.0
                        result['validation_passed'] = True
                        result['validation_reason'] = "LLM validation disabled"
                        result['validation_method'] = 'disabled'
                        status = "✅ FOUND CONTENT" if bank_found else "✅ NO CONTENT"
                    else:
                        # No content found - no validation needed
                        result['validation_score'] = 1.0
                        result['validation_passed'] = True
                        result['validation_reason'] = "No content to validate"
                        result['validation_method'] = 'not_applicable'
                        status = "✅ NO CONTENT"
                    
                    tracker.add_result(bank_code, script_name, display_name, result)
                    
                    if bank_found:
                        content_found_count += 1
                        print(f"   📊 Content details: {len(result.get('relevant_pages', []))} pages, {len(result.get('highlights', []))} highlights")
                
                print(f"   {status}: {bank_name}")
                
                # Update frontend progress after EACH bank with specific message
                if task_id in active_tasks:
                    current_progress = active_tasks[task_id]["progress"]["current"]
                    new_progress = current_progress + 1
                    
                    active_tasks[task_id]["progress"]["current"] = new_progress
                    active_tasks[task_id]["progress"]["current_task"] = f"✅ Completed {bank_name} in {display_name} ({new_progress}/{tracker.total_combinations})"
                    
                    # Update results immediately for real-time frontend sync
                    active_tasks[task_id]["results"] = format_results_for_frontend(tracker, selected_banks, [script_name])
                    active_tasks[task_id]["tracker_stats"] = tracker.get_summary_stats()
                    
                    print(f"📡 WebSocket update: Bank {bank_name} completed ({new_progress}/{tracker.total_combinations})")
                    await asyncio.sleep(0.3)
                    
            else:
                tracker.add_result(bank_code, script_name, display_name, None)
                bank_name = get_bank_info(bank_code)['english_name']
                print(f"   ❌ FAILED: {bank_name}")
                
                # Update frontend progress even for failures
                if task_id in active_tasks:
                    current_progress = active_tasks[task_id]["progress"]["current"]
                    new_progress = current_progress + 1
                    
                    active_tasks[task_id]["progress"]["current"] = new_progress
                    active_tasks[task_id]["progress"]["current_task"] = f"❌ Failed {bank_name} in {display_name} ({new_progress}/{tracker.total_combinations})"
                    
                    active_tasks[task_id]["results"] = format_results_for_frontend(tracker, selected_banks, [script_name])
                    active_tasks[task_id]["tracker_stats"] = tracker.get_summary_stats()
                    
                    print(f"📡 WebSocket update: Bank {bank_name} failed ({new_progress}/{tracker.total_combinations})")
                    await asyncio.sleep(0.3)
        
        # Send completion update for this newspaper
        if task_id in active_tasks:
            validation_summary = ""
            if llm_validation_enabled:
                validation_summary = " (with LLM validation)"
            
            active_tasks[task_id]["progress"].update({
                "current_task": f"🎯 Completed {display_name}{validation_summary} - {content_found_count} banks with content",
                "phase": "analyzing"
            })
            print(f"📡 WebSocket update: {display_name} completed")
            await asyncio.sleep(0.5)
        
        print(f"📊 {display_name} analysis summary: {content_found_count} banks with LLM-validated content found")
        print(f"✅ {display_name} analysis completed")
        
    except Exception as e:
        error_msg = f"Error processing {display_name}: {str(e)}"
        print(f"❌ {error_msg}")
        import traceback
        traceback.print_exc()
        
        # Add failed results for all banks
        for bank_code in selected_banks:
            tracker.add_result(bank_code, script_name, display_name, None)
            
            # Update frontend progress even for exceptions
            if task_id in active_tasks:
                current_progress = active_tasks[task_id]["progress"]["current"]
                active_tasks[task_id]["progress"]["current"] = current_progress + 1
                active_tasks[task_id]["progress"]["current_task"] = f"❌ Error in {display_name}: {str(e)}"
                active_tasks[task_id]["results"] = format_results_for_frontend(tracker, selected_banks, [script_name])
                active_tasks[task_id]["tracker_stats"] = tracker.get_summary_stats()
                await asyncio.sleep(0.3)
                
async def cleanup_temp_files(tracker):
    """ENHANCED: Clean up temporary files AFTER reports are generated"""
    print("\n🗂️ Cleaning up temporary files AFTER report generation...")
    
    screenshot_files_removed = 0
    pdf_files_removed = 0
    chart_files_removed = 0
    
    # Method 1: Clean up via tracker results (existing code...)
    for bank_code in tracker.bank_stats.keys():
        bank_results = tracker.get_bank_results(bank_code)
        
        for script_name, result_data in bank_results.items():
            result = result_data['result']
            if result:
                # Clean up screenshots from results (existing code...)
                if result.get('page_screenshots'):
                    try:
                        for screenshot_path in result['page_screenshots'].values():
                            if os.path.exists(screenshot_path):
                                os.remove(screenshot_path)
                                screenshot_files_removed += 1
                                print(f"   🗂️ Removed screenshot: {os.path.basename(screenshot_path)}")
                            
                            # Also remove processed versions
                            processed_path = screenshot_path.replace('.png', '_cropped.png')
                            if os.path.exists(processed_path):
                                os.remove(processed_path)
                                screenshot_files_removed += 1
                                print(f"   🗂️ Removed processed screenshot: {os.path.basename(processed_path)}")
                    
                    except Exception as e:
                        print(f"   ⚠️ Warning cleaning up screenshots: {e}")
                
                # Clean up PDF files (existing code...)
                if result.get('pdf_path'):
                    try:
                        pdf_path = result['pdf_path']
                        if os.path.exists(pdf_path):
                            os.remove(pdf_path)
                            pdf_files_removed += 1
                            print(f"   🗂️ Removed PDF: {os.path.basename(pdf_path)}")
                    except Exception as e:
                        print(f"   ⚠️ Warning cleaning up PDF: {e}")
    
    # Method 2: Clean up chart files
    try:
        charts_dir = "Charts"
        if os.path.exists(charts_dir):
            for filename in os.listdir(charts_dir):
                if filename.endswith('.png'):
                    file_path = os.path.join(charts_dir, filename)
                    try:
                        os.remove(file_path)
                        chart_files_removed += 1
                        print(f"   🗂️ Removed chart: {filename}")
                    except Exception as e:
                        print(f"   ⚠️ Could not remove chart {filename}: {e}")
    except Exception as e:
        print(f"   ⚠️ Warning cleaning up charts: {e}")
    
    # Method 3: Fallback - scan current directory for temp files (existing code...)
    try:
        current_dir = os.getcwd()
        print(f"   🔍 Scanning {current_dir} for remaining temp files...")
        
        for filename in os.listdir(current_dir):
            # Look for temp screenshot patterns
            if (filename.startswith('temp_page_') and filename.endswith('.png')) or \
               (filename.startswith('AlRai_Media_') and filename.endswith('.pdf')):
                
                file_path = os.path.join(current_dir, filename)
                try:
                    os.remove(file_path)
                    if filename.endswith('.png'):
                        screenshot_files_removed += 1
                        print(f"   🗂️ Removed orphaned screenshot: {filename}")
                    else:
                        pdf_files_removed += 1
                        print(f"   🗂️ Removed orphaned PDF: {filename}")
                except Exception as e:
                    print(f"   ⚠️ Could not remove {filename}: {e}")
                    
    except Exception as e:
        print(f"   ⚠️ Warning scanning directory for cleanup: {e}")
    
    print(f"🗂️ Cleanup completed: {screenshot_files_removed} screenshots, {pdf_files_removed} PDFs, {chart_files_removed} charts removed")
    

# Also add this function to manually clean up any remaining files
def manual_cleanup_temp_files():
    """Manual cleanup function to remove any remaining temp files"""
    print("🗂️ Manual cleanup of temp files...")
    
    current_dir = os.getcwd()
    removed_count = 0
    
    try:
        for filename in os.listdir(current_dir):
            if (filename.startswith('temp_page_') and filename.endswith('.png')) or \
               (filename.startswith('AlRai_Media_') and filename.endswith('.pdf')):
                
                file_path = os.path.join(current_dir, filename)
                try:
                    os.remove(file_path)
                    removed_count += 1
                    print(f"   🗂️ Removed: {filename}")
                except Exception as e:
                    print(f"   ❌ Could not remove {filename}: {e}")
    
    except Exception as e:
        print(f"❌ Error during manual cleanup: {e}")
    
    print(f"🗂️ Manual cleanup completed: {removed_count} files removed")
    return removed_count
# ==================================================================================
# FIXED FASTAPI ENDPOINTS WITH ENHANCED FRONTEND SYNC
# ==================================================================================

@app.get("/")
async def root():
    return {"message": "Kuwait Banking News Analyzer API v2.2 - FIXED", "status": "running", "methodology": "master_orchestrator"}

@app.get("/api/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now(KW).isoformat(),
        "active_tasks": len(active_tasks),
        "master_orchestrator": "integrated",
        "asyncio_fix": "applied",
        "available_banks": len(BANK_CONFIGS),
        "available_newspapers": len(NEWSPAPER_SCRIPTS),
        "version": "2.2.0-FIXED"
    }

@app.post("/api/start-analysis")
async def start_analysis(request: AnalysisRequest, background_tasks: BackgroundTasks):
    """FIXED: Start a new banking news analysis using master orchestrator methodology"""
    
    # Validate inputs against available configurations
    valid_banks = [bank for bank in request.selected_banks if bank in BANK_CONFIGS]
    valid_newspapers = [news for news in request.selected_newspapers if news in NEWSPAPER_SCRIPTS]
    
    if not valid_banks:
        raise HTTPException(status_code=400, detail="No valid banks selected")
    if not valid_newspapers:
        raise HTTPException(status_code=400, detail="No valid newspapers selected")
    
    task_id = str(uuid.uuid4())
    total_combinations = len(valid_banks) * len(valid_newspapers)
    
    valid_models = ["gpt-4o", "gpt-4o-mini"]
    selected_model = request.selected_model if request.selected_model in valid_models else "gpt-4o"
    
    # Initialize task with master orchestrator structure
    active_tasks[task_id] = {
        "status": "starting",
        "progress": {
            "current": 0,
            "total": total_combinations,
            "phase": "idle",
            "current_task": "Initializing master orchestrator analysis..."
        },
        "results": {},
        "tracker": None,  # Will hold MultiBankAnalysisTracker instance
        "config": {
            "banks": valid_banks,
            "newspapers": valid_newspapers,
            "model": selected_model
        },
        "created_at": datetime.now(KW).isoformat(),
        "error": None,
        "tracker_stats": {}
    }
    
    # Start background analysis using master orchestrator
    background_tasks.add_task(run_master_orchestrator_analysis, task_id, valid_banks, valid_newspapers)
    
    return {
        "task_id": task_id, 
        "status": "started", 
        "total_combinations": total_combinations,
        "methodology": "master_orchestrator",
        "banks": [get_bank_info(bank)['english_name'] for bank in valid_banks],
        "newspapers": [NEWSPAPER_SCRIPTS[news] for news in valid_newspapers]
    }

async def run_master_orchestrator_analysis(task_id: str, selected_banks: List[str], selected_newspapers: List[str], selected_model: str = "gpt-4o"):
    """FIXED: Run the analysis using the master orchestrator methodology with enhanced frontend sync"""
    
    try:
        print(f"🚀 Starting Master Orchestrator Analysis for task {task_id}")
        print(f"📊 Banks: {', '.join([get_bank_info(bank)['english_name'] for bank in selected_banks])}")
        print(f"📰 Newspapers: {', '.join([NEWSPAPER_SCRIPTS[news] for news in selected_newspapers])}")
        print("=" * 80)
        
        # Initialize the MultiBankAnalysisTracker
        tracker = MultiBankAnalysisTracker()
        tracker.total_combinations = len(selected_banks) * len(selected_newspapers)
        
        # Store tracker in task
        active_tasks[task_id]["tracker"] = tracker
        active_tasks[task_id]["status"] = "running"
        
        # Update progress
        active_tasks[task_id]["progress"].update({
            "current": 0,
            "total": tracker.total_combinations,
            "phase": "analyzing",
            "current_task": "Starting newspaper analysis with master orchestrator..."
        })
        
        # Process each newspaper using master orchestrator methodology
        for i, newspaper_id in enumerate(selected_newspapers):
            if task_id not in active_tasks:  # Task was cancelled
                break
                
            newspaper_name = NEWSPAPER_SCRIPTS.get(newspaper_id, newspaper_id)
            
            print(f"\n📰 Processing {newspaper_name} ({i+1}/{len(selected_newspapers)})...")
            
            try:
                # Use enhanced load_and_run_newspaper_script function with frontend sync
                await load_and_run_newspaper_script(newspaper_id, newspaper_name, tracker, selected_banks, task_id)
                
                print(f"✅ Completed {newspaper_name} analysis")
                
                # Update progress after each newspaper with enhanced stats
                completed_so_far = (i + 1) * len(selected_banks)
                if task_id in active_tasks:
                    active_tasks[task_id]["progress"].update({
                        "current": completed_so_far,
                        "current_task": f"Completed {newspaper_name} analysis ({completed_so_far}/{tracker.total_combinations})",
                        "phase": "analyzing"
                    })
                    
                    # Update tracker stats for frontend
                    active_tasks[task_id]["tracker_stats"] = tracker.get_summary_stats()
                    # Update results with all newspapers processed so far
                    active_tasks[task_id]["results"] = format_results_for_frontend(tracker, selected_banks, selected_newspapers[:i+1])
                
            except Exception as e:
                print(f"❌ Error processing {newspaper_name}: {str(e)}")
                
                # Add failed results for all banks for this newspaper
                for bank_code in selected_banks:
                    tracker.add_result(bank_code, newspaper_id, newspaper_name, None)
                    
                    # Update frontend progress even on failure
                    if task_id in active_tasks:
                        current_progress = active_tasks[task_id]["progress"]["current"]
                        active_tasks[task_id]["progress"]["current"] = current_progress + 1
                        active_tasks[task_id]["results"] = format_results_for_frontend(tracker, selected_banks, selected_newspapers)
                        active_tasks[task_id]["tracker_stats"] = tracker.get_summary_stats()
        
        if task_id in active_tasks:
            # Generate reports using master orchestrator methodology
            active_tasks[task_id]["progress"].update({
                "current": tracker.total_combinations,
                "current_task": "Generating comprehensive reports using master orchestrator...",
                "phase": "generating"
            })
            
            await asyncio.sleep(2)
            
            print("\n📄 Generating reports using master orchestrator methodology...")
            
            today = datetime.now(KW)
            report_paths = {}
            
            try:
                # Generate individual bank reports using master orchestrator
                bank_reports = []
                for j, bank_code in enumerate(selected_banks):
                    try:
                        # Update progress for each bank report
                        bank_name = get_bank_info(bank_code)['english_name']
                        active_tasks[task_id]["progress"].update({
                            "current_task": f"Generating report for {bank_name} ({j+1}/{len(selected_banks)})...",
                            "phase": "generating"
                        })
                        
                        print(f"📊 Generating report for {bank_name}...")
                        report_path = create_individual_bank_report(bank_code, tracker, today)
                        bank_reports.append(report_path)
                        print(f"✅ Generated report for {bank_name}")

                    except Exception as e:
                        print(f"❌ Error generating report for {get_bank_info(bank_code)['english_name']}: {e}")
               
                # Generate consolidated headlines report using master orchestrator
                # Generate consolidated headlines report using master orchestrator
                try:
                    active_tasks[task_id]["progress"].update({
                        "current_task": "Generating consolidated headlines report...",
                        "phase": "generating"
                    })
                    
                    print("📑 Generating consolidated headlines report...")
                    headlines_report_path = create_headlines_report(tracker, today)
                    print("✅ Generated consolidated headlines report")
                except Exception as e:
                    print(f"❌ Error generating headlines report: {e}")
                
                # Generate analytics report with charts
                try:
                    active_tasks[task_id]["progress"].update({
                        "current_task": "Generating analytics report with charts...",
                        "phase": "generating"
                    })
                    
                    print("📊 Generating analytics report...")
                    analytics_report_path = create_analytics_report(tracker, today)
                    print("✅ Generated analytics report")
                except Exception as e:
                    print(f"❌ Error generating analytics report: {e}")
                    analytics_report_path = None
                # try:
                #     active_tasks[task_id]["progress"].update({
                #         "current_task": "Generating consolidated headlines report...",
                #         "phase": "generating"
                #     })
                    
                #     print("📑 Generating consolidated headlines report...")
                #     headlines_report_path = create_headlines_report(tracker, today)
                #     print("✅ Generated consolidated headlines report")
                # except Exception as e:
                #     print(f"❌ Error generating headlines report: {e}")
                report_paths = {
                    "bank_reports": bank_reports,
                    "headlines_report": headlines_report_path,
                    "analytics_report": analytics_report_path
                    }
                # report_paths = {
                #     "bank_reports": bank_reports,
                #     "headlines_report": headlines_report_path
                # }
                
                # CRITICAL FIX: Ensure final status is properly set with download links
                if task_id in active_tasks:
                    # Store report paths
                    active_tasks[task_id]["report_paths"] = report_paths
                    
                    # FORCE final status update
                    active_tasks[task_id]["status"] = "completed"
                    active_tasks[task_id]["progress"].update({
                        "current": tracker.total_combinations,
                        "current_task": "Master orchestrator analysis complete!",
                        "phase": "complete"
                    })
                    
                    # CRITICAL: Force immediate sync of all data
                    active_tasks[task_id]["tracker_stats"] = tracker.get_summary_stats()
                    active_tasks[task_id]["results"] = format_results_for_frontend(tracker, selected_banks, selected_newspapers)
                    
                    # Add explicit completion markers
                    active_tasks[task_id]["analysis_complete"] = True
                    active_tasks[task_id]["reports_ready"] = True
                    active_tasks[task_id]["completion_time"] = datetime.now(KW).isoformat()
                    
                    print(f"⚙️ FORCING FINAL STATUS UPDATE for task {task_id}")
                    print(f"   Status: {active_tasks[task_id]['status']}")
                    print(f"   Reports ready: {bool(active_tasks[task_id].get('report_paths'))}")
                    print(f"   Bank reports: {len(report_paths.get('bank_reports', []))}")
                    print(f"   Headlines report: {bool(report_paths.get('headlines_report'))}")
                    
                    # Give WebSocket time to pick up the changes
                    await asyncio.sleep(3)  # Longer delay for WebSocket sync
                    
                    print(f"⚙️ Final status set for task {task_id} - WebSocket should pick this up")
                
                print("📊 Report generation completed:")
                print(f"📑 Headlines report: {headlines_report_path}")
                print(f"🏦 Bank reports: {len(bank_reports)} generated")
                
            except Exception as e:
                print(f"❌ Error in report generation: {e}")
                import traceback
                traceback.print_exc()
            
            # Clean up temporary files using master orchestrator
            try:
                active_tasks[task_id]["progress"].update({
                    "current_task": "Cleaning up temporary files...",
                    "phase": "generating"
                })
                
                await cleanup_temp_files(tracker)
                print("🧹 Cleanup completed")
            except Exception as e:
                print(f"⚠️ Warning during cleanup: {e}")

            # Print comprehensive summary using master orchestrator
            print_master_orchestrator_summary(tracker, selected_banks, selected_newspapers)
            print(f"✅ Master orchestrator analysis task {task_id} completed successfully")
        
    except Exception as e:
        print(f"💥 Critical error in master orchestrator analysis task {task_id}: {str(e)}")
        import traceback
        traceback.print_exc()
        
        if task_id in active_tasks:
            active_tasks[task_id]["status"] = "error"
            active_tasks[task_id]["error"] = str(e)
            active_tasks[task_id]["progress"]["phase"] = "error"
            active_tasks[task_id]["progress"]["current_task"] = f"Master orchestrator error: {str(e)}"

def print_master_orchestrator_summary(tracker: MultiBankAnalysisTracker, selected_banks: List[str], selected_newspapers: List[str]):
    """Print comprehensive summary using master orchestrator methodology"""
    
    stats = tracker.get_summary_stats()
    print("\n" + "=" * 80)
    print("📊 MASTER ORCHESTRATOR ANALYSIS SUMMARY")
    print("=" * 80)
    print(f"📄 Total combinations analyzed: {stats['completed_combinations']}/{stats['total_combinations']}")
    print(f"✅ Combinations with content: {stats['combinations_with_content']}")
    print(f"🏦 Banks analyzed: {stats['banks_analyzed']}")
    print(f"📰 Newspapers analyzed: {stats['newspapers_analyzed']}")
    
    
    # Show results by bank
    print(f"\n📊 Results by Bank:")
    for bank_code in selected_banks:
        bank_stats = tracker.bank_stats.get(bank_code, {})
        bank_name = bank_stats.get('name', bank_code)
        content_count = bank_stats.get('newspapers_with_content', 0)
        total_analyzed = bank_stats.get('newspapers_analyzed', 0)
        pages_found = bank_stats.get('total_pages_found', 0)
        print(f"   🏦 {bank_name}: {content_count}/{total_analyzed} newspapers, {pages_found} pages")
    
    # Show results by newspaper
    print(f"\n📊 Results by Newspaper:")
    for script_name, newspaper_data in tracker.newspaper_stats.items():
        name = newspaper_data['name']
        content_count = newspaper_data['banks_with_content']
        total_analyzed = newspaper_data['banks_analyzed']
        pages_found = newspaper_data['total_pages_found']
        
        print(f"   📰 {name}: {content_count}/{total_analyzed} banks, {pages_found} pages")
    
    print("=" * 80)

# FIXED WebSocket handler
# WEBSOCKET TIMEOUT FIX FOR LARGE ANALYSIS
# Replace the websocket_endpoint function in main.py

@app.websocket("/ws/analysis/{task_id}")
async def websocket_endpoint(websocket: WebSocket, task_id: str):
    """ENHANCED WebSocket with aggressive connection management for download UI"""
    await websocket.accept()
    
    try:
        print(f"🔌 WebSocket CONNECTED for task {task_id}")
        
        # Send initial connection confirmation
        initial_message = {
            "status": "connected",
            "message": "WebSocket established - monitoring for completion",
            "task_id": task_id,
            "timestamp": datetime.now(KW).isoformat()
        }
        
        try:
            await websocket.send_text(json.dumps(initial_message))
            print(f"📤 Initial message sent for task {task_id}")
        except Exception as e:
            print(f"❌ Failed to send initial message: {e}")
            return
        
        last_update = None
        heartbeat_count = 0
        update_count = 0
        max_connection_time = 1800  # 30 minutes max connection
        connection_start = datetime.now(KW)
        
        while True:
            try:
                # Check connection time
                connection_duration = (datetime.now(KW) - connection_start).total_seconds()
                if connection_duration > max_connection_time:
                    print(f"⏰ WebSocket connection timeout after {connection_duration}s for task {task_id}")
                    break
                
                if task_id in active_tasks:
                    task = active_tasks[task_id]
                    
                    # Get tracker statistics
                    tracker_stats = {}
                    try:
                        if task.get("tracker"):
                            tracker_stats = task["tracker"].get_summary_stats()
                    except Exception as e:
                        print(f"⚠️ Error getting tracker stats: {e}")
                        tracker_stats = {}
                    
                    # Prepare current update
                    current_update = {
                        "task_id": task_id,
                        "status": task.get("status", "unknown"),
                        "progress": task.get("progress", {}),
                        "results": task.get("results", {}),
                        "error": task.get("error"),
                        "report_paths": task.get("report_paths"),
                        "tracker_stats": tracker_stats,
                        "methodology": "master_orchestrator",
                        "config": task.get("config", {}),
                        "timestamp": datetime.now(KW).isoformat(),
                        "connection_duration": int(connection_duration)
                    }
                    
                    # CRITICAL: Always send updates for completed tasks
                    task_status = task.get("status")
                    is_completed = task_status in ["completed", "error"]
                    has_reports = bool(task.get("report_paths"))
                    data_changed = current_update != last_update
                    
                    if data_changed or is_completed:
                        try:
                            await websocket.send_text(json.dumps(current_update))
                            last_update = current_update.copy()
                            update_count += 1
                            heartbeat_count = 0
                            
                            print(f"📤 Update #{update_count} sent for task {task_id}")
                            print(f"   Status: {task_status}")
                            print(f"   Has reports: {has_reports}")
                            
                            if is_completed and has_reports:
                                print(f"🎯 COMPLETED TASK WITH REPORTS - sending final confirmation")
                                
                                # Send additional confirmation
                                confirmation = {
                                    "task_id": task_id,
                                    "status": "reports_available",
                                    "message": "Analysis complete - reports ready for download",
                                    "report_paths": task.get("report_paths"),
                                    "timestamp": datetime.now(KW).isoformat(),
                                    "final_confirmation": True
                                }
                                
                                await asyncio.sleep(1)  # Brief pause
                                await websocket.send_text(json.dumps(confirmation))
                                print(f"📤 Final confirmation sent for task {task_id}")

                                # Stay connected longer for completed tasks
                                print(f"⏳ Holding connection for 10 seconds to ensure frontend receives updates...")
                                await asyncio.sleep(10)
                                print(f"✅ Connection hold complete for task {task_id}")
                                break
                        
                        except Exception as e:
                            print(f"❌ Error sending update for task {task_id}: {e}")
                            break
                    
                    else:
                        # Send heartbeat
                        heartbeat_count += 1
                        if heartbeat_count >= 5:  # Every 5 seconds
                            try:
                                heartbeat = {
                                    "status": "heartbeat",
                                    "task_id": task_id,
                                    "message": f"Analysis in progress... ({int(connection_duration)}s)",
                                    "current_progress": task.get("progress", {}).get("current", 0),
                                    "total_progress": task.get("progress", {}).get("total", 0),
                                    "timestamp": datetime.now(KW).isoformat()
                                }
                                await websocket.send_text(json.dumps(heartbeat))
                                heartbeat_count = 0
                                print(f"💓 Heartbeat sent for task {task_id} ({int(connection_duration)}s)")
                            except Exception as e:
                                print(f"❌ Error sending heartbeat: {e}")
                                
                    
                    # Break if task completed and we've sent the final updates
                    if is_completed and has_reports and update_count > 0:
                        print(f"🏁 Task {task_id} completed with reports - ending WebSocket connection")
                        break
                
                else:
                    # Task not found
                    print(f"❌ Task {task_id} not found in active_tasks")
                    try:
                        not_found = {
                            "task_id": task_id,
                            "status": "not_found",
                            "error": "Task not found or was cancelled",
                            "timestamp": datetime.now(KW).isoformat()
                        }
                        await websocket.send_text(json.dumps(not_found))
                    except Exception as e:
                        print(f"❌ Error sending not found message: {e}")
                    break
                
                # Shorter sleep for more responsive updates
                await asyncio.sleep(1)
                
            except Exception as loop_error:
                print(f"❌ WebSocket loop error for task {task_id}: {loop_error}")
                break
            
    except Exception as e:
        print(f"❌ WebSocket error for task {task_id}: {e}")
        import traceback
        traceback.print_exc()
    
    finally:
        print(f"🔌 WebSocket DISCONNECTED for task {task_id} after {update_count} updates")

# FIXED: Enhanced task status endpoint
@app.get("/api/progress/{task_id}")
async def get_progress(task_id: str):
    """Get current progress for a specific task with enhanced error handling"""
    try:
        if task_id not in active_tasks:
            raise HTTPException(status_code=404, detail="Task not found")
        
        task = active_tasks[task_id]
        
        # Include tracker statistics if available with error handling
        tracker_stats = {}
        try:
            if task.get("tracker"):
                tracker_stats = task["tracker"].get_summary_stats()
        except Exception as e:
            print(f"⚠️ Warning: Could not get tracker stats: {e}")
            tracker_stats = {}
        
        # Enhanced response with all needed frontend data
        response = {
            "task_id": task_id,
            "status": task.get("status", "unknown"),
            "progress": task.get("progress", {
                "current": 0,
                "total": 0,
                "phase": "idle",
                "current_task": "Initializing..."
            }),
            "results": task.get("results", {}),
            "error": task.get("error"),
            "report_paths": task.get("report_paths"),
            "tracker_stats": tracker_stats,
            "methodology": "master_orchestrator",
            "config": task.get("config", {}),
            "created_at": task.get("created_at"),
            "timestamp": datetime.now(KW).isoformat()
        }
        
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error getting progress for task {task_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving task progress: {str(e)}")

@app.get("/api/download-report/{task_id}/{report_type}")
async def download_individual_report(task_id: str, report_type: str):
    """Download individual Word document reports generated by master orchestrator"""
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Analysis not completed")
    
    report_paths = task.get("report_paths")
    if not report_paths:
        raise HTTPException(status_code=404, detail="No reports generated by master orchestrator")
    
    try:
        if report_type == "headlines":
            headlines_report = report_paths.get("headlines_report")
            if not headlines_report or not os.path.exists(headlines_report):
                raise HTTPException(status_code=404, detail="Master orchestrator headlines report not found")
            
            return FileResponse(
                headlines_report,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=os.path.basename(headlines_report),
                headers={"Content-Disposition": f"attachment; filename={os.path.basename(headlines_report)}"}
            )
        
        elif report_type.startswith("bank_"):
            bank_code = report_type.replace("bank_", "")
            bank_reports = report_paths.get("bank_reports", [])
            
            bank_info = get_bank_info(bank_code)
            bank_name = bank_info['english_name'] if bank_info else bank_code
            expected_filename_pattern = bank_name.replace(' ', '_')
            
            matching_report = None
            for report_path in bank_reports:
                if expected_filename_pattern in os.path.basename(report_path) and os.path.exists(report_path):
                    matching_report = report_path
                    break
            
            if not matching_report:
                raise HTTPException(status_code=404, detail=f"Master orchestrator report for {bank_name} not found")
            
            return FileResponse(
                matching_report,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=os.path.basename(matching_report),
                headers={"Content-Disposition": f"attachment; filename={os.path.basename(matching_report)}"}
            )
        
        else:
            raise HTTPException(status_code=400, detail="Invalid report type")
            
    except Exception as e:
        print(f"❌ Error downloading master orchestrator report: {e}")
        raise HTTPException(status_code=500, detail=f"Error downloading report: {str(e)}")

@app.post("/api/download-reports/{task_id}")
async def download_reports(task_id: str):
    """Download analysis reports as a ZIP file containing Word documents"""
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Analysis not completed")
    
    try:
        # Create zip file with all Word document reports
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            # Check if we have report paths from the master orchestrator
            if 'report_paths' in task:
                report_paths = task['report_paths']
                
                # Add individual bank reports
                for report_path in report_paths.get('bank_reports', []):
                    if os.path.exists(report_path):
                        zip_file.write(report_path, os.path.basename(report_path))
                        print(f"📄 Added to zip: {os.path.basename(report_path)}")
                
                # Add consolidated headlines report
                headlines_report = report_paths.get('headlines_report')
                if headlines_report and os.path.exists(headlines_report):
                    zip_file.write(headlines_report, os.path.basename(headlines_report))
                    print(f"📄 Added to zip: {os.path.basename(headlines_report)}")
        
        zip_buffer.seek(0)
        
        # Create filename with timestamp
        timestamp = datetime.now(KW).strftime('%Y%m%d_%H%M%S')
        filename = f"master_orchestrator_banking_analysis_{timestamp}.zip"
        
        return StreamingResponse(
            io.BytesIO(zip_buffer.read()),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        print(f"âŒ Error generating reports for task {task_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating reports: {str(e)}")

@app.get("/api/download-analytics/{task_id}")
async def download_analytics_report(task_id: str):
    """Download analytics report with charts"""
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Analysis not completed")
    
    report_paths = task.get("report_paths")
    if not report_paths:
        raise HTTPException(status_code=404, detail="No reports generated")
    
    analytics_report = report_paths.get("analytics_report")
    if not analytics_report or not os.path.exists(analytics_report):
        raise HTTPException(status_code=404, detail="Analytics report not found")
    
    return FileResponse(
        analytics_report,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(analytics_report),
        headers={"Content-Disposition": f"attachment; filename={os.path.basename(analytics_report)}"}
    )

@app.get("/api/download-report/{task_id}/{report_type}")
async def download_individual_report(task_id: str, report_type: str):
    """Download individual Word document reports generated by master orchestrator"""
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Analysis not completed")
    
    report_paths = task.get("report_paths")
    if not report_paths:
        raise HTTPException(status_code=404, detail="No reports generated by master orchestrator")
    
    try:
        if report_type == "headlines":
            headlines_report = report_paths.get("headlines_report")
            if not headlines_report or not os.path.exists(headlines_report):
                raise HTTPException(status_code=404, detail="Master orchestrator headlines report not found")
            
            return FileResponse(
                headlines_report,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=os.path.basename(headlines_report),
                headers={"Content-Disposition": f"attachment; filename={os.path.basename(headlines_report)}"}
            )
        
        elif report_type.startswith("bank_"):
            bank_code = report_type.replace("bank_", "")
            bank_reports = report_paths.get("bank_reports", [])
            
            bank_info = get_bank_info(bank_code)
            bank_name = bank_info['english_name'] if bank_info else bank_code
            expected_filename_pattern = bank_name.replace(' ', '_')
            
            matching_report = None
            for report_path in bank_reports:
                if expected_filename_pattern in os.path.basename(report_path) and os.path.exists(report_path):
                    matching_report = report_path
                    break
            
            if not matching_report:
                raise HTTPException(status_code=404, detail=f"Master orchestrator report for {bank_name} not found")
            
            return FileResponse(
                matching_report,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=os.path.basename(matching_report),
                headers={"Content-Disposition": f"attachment; filename={os.path.basename(matching_report)}"}
            )
        
        else:
            raise HTTPException(status_code=400, detail="Invalid report type")
            
    except Exception as e:
        print(f"❌ Error downloading master orchestrator report: {e}")
        raise HTTPException(status_code=500, detail=f"Error downloading report: {str(e)}")


@app.post("/api/cleanup-temp-files")
async def manual_cleanup_endpoint():
    """Manual cleanup endpoint to remove any remaining temp files"""
    try:
        removed_count = manual_cleanup_temp_files()
        return {
            "status": "success",
            "message": f"Manual cleanup completed: {removed_count} files removed",
            "removed_count": removed_count
        }
    except Exception as e:
        return {
            "status": "error", 
            "message": f"Manual cleanup failed: {str(e)}"
        }

# Also add an endpoint to check what tasks exist:
@app.get("/api/debug/active-tasks")
async def debug_active_tasks():
    """Debug endpoint to see what tasks are currently active"""
    return {
        "active_task_count": len(active_tasks),
        "active_task_ids": list(active_tasks.keys()),
        "task_details": {
            task_id: {
                "status": task_data.get("status"),
                "created_at": task_data.get("created_at"),
                "phase": task_data.get("progress", {}).get("phase"),
                "has_reports": bool(task_data.get("report_paths"))
            }
            for task_id, task_data in active_tasks.items()
        }
    }

@app.get("/api/list-reports/{task_id}")
async def list_available_reports(task_id: str):
    """List all available reports for a completed task"""
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    if task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Analysis not completed")
    
    report_paths = task.get("report_paths")
    if not report_paths:
        return {"reports": []}
    
    reports = []
    
    # 1. Add headlines report FIRST
    headlines_report = report_paths.get("headlines_report")
    if headlines_report and os.path.exists(headlines_report):
        reports.append({
            "type": "headlines",
            "name": "Consolidated Headlines Report",
            "filename": os.path.basename(headlines_report),
            "download_url": f"/api/download-report/{task_id}/headlines"
        })
    
    # 2. Add analytics report SECOND
    analytics_report = report_paths.get("analytics_report")
    if analytics_report and os.path.exists(analytics_report):
        reports.append({
            "type": "analytics",
            "name": "Analytics Report with Charts",
            "filename": os.path.basename(analytics_report),
            "download_url": f"/api/download-analytics/{task_id}"
        })
    
    # 3. Add bank reports LAST
    for report_path in report_paths.get("bank_reports", []):
        if os.path.exists(report_path):
            filename = os.path.basename(report_path)
            for bank_code in BANK_CONFIGS.keys():
                bank_info = get_bank_info(bank_code)
                bank_name = bank_info['english_name'].replace(' ', '_')
                if bank_name in filename:
                    reports.append({
                        "type": f"bank_{bank_code}",
                        "name": f"{bank_info['english_name']} Report",
                        "filename": filename,
                        "download_url": f"/api/download-report/{task_id}/bank_{bank_code}"
                    })
                    break
    
    return {"reports": reports}

# @app.get("/api/list-reports/{task_id}")
# async def list_available_reports(task_id: str):
#     """List all available reports for a completed task"""
#     if task_id not in active_tasks:
#         raise HTTPException(status_code=404, detail="Task not found")
    
#     task = active_tasks[task_id]
#     if task["status"] != "completed":
#         raise HTTPException(status_code=400, detail="Analysis not completed")
    
#     report_paths = task.get("report_paths")
#     if not report_paths:
#         return {"reports": []}
    
#     reports = []
    
#     # Add headlines report
#     headlines_report = report_paths.get("headlines_report")
#     if headlines_report and os.path.exists(headlines_report):
#         reports.append({
#             "type": "headlines",
#             "name": "Consolidated Headlines Report",
#             "filename": os.path.basename(headlines_report),
#             "download_url": f"/api/download-report/{task_id}/headlines"
#         })
    
#     # Add bank reports
#     for report_path in report_paths.get("bank_reports", []):
#         if os.path.exists(report_path):
#             filename = os.path.basename(report_path)
#             for bank_code in BANK_CONFIGS.keys():
#                 bank_info = get_bank_info(bank_code)
#                 bank_name = bank_info['english_name'].replace(' ', '_')
#                 if bank_name in filename:
#                     reports.append({
#                         "type": f"bank_{bank_code}",
#                         "name": f"{bank_info['english_name']} Report",
#                         "filename": filename,
#                         "download_url": f"/api/download-report/{task_id}/bank_{bank_code}"
#                     })
#                     break
    
#     # Add analytics report
#     analytics_report = report_paths.get("analytics_report")
#     if analytics_report and os.path.exists(analytics_report):
#         reports.append({
#             "type": "analytics",
#             "name": "Analytics Report with Charts",
#             "filename": os.path.basename(analytics_report),
#             "download_url": f"/api/download-analytics/{task_id}"
#         })
    
#     return {"reports": reports}

@app.delete("/api/tasks/{task_id}")
async def cancel_task(task_id: str):
    """Cancel a running task"""
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    # Remove the task (this will cause the background task to stop)
    del active_tasks[task_id]
    
    return {"message": f"Master orchestrator task {task_id} cancelled"}

@app.get("/api/tasks")
async def list_tasks():
    """List all active tasks"""
    return {
        "active_tasks": len(active_tasks),
        "methodology": "master_orchestrator",
        "tasks": [
            {
                "task_id": task_id,
                "status": task_data["status"],
                "created_at": task_data.get("created_at"),
                "progress": task_data["progress"],
                "tracker_stats": task_data.get("tracker_stats", {}),
                "config": task_data.get("config", {})
            }
            for task_id, task_data in active_tasks.items()
        ]
    }

# ADD THIS NEW ENDPOINT TO YOUR MAIN.PY (after the other @app.get endpoints):

@app.get("/api/task-status/{task_id}")
async def get_task_status_with_reports(task_id: str):
    """Backup endpoint to check task status and report availability"""
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = active_tasks[task_id]
    
    # Enhanced response with all needed data
    response = {
        "task_id": task_id,
        "status": task.get("status", "unknown"),
        "progress": task.get("progress", {}),
        "results": task.get("results", {}),
        "error": task.get("error"),
        "report_paths": task.get("report_paths"),
        "reports_available": bool(task.get("report_paths")),
        "analysis_complete": task.get("analysis_complete", False),
        "reports_ready": task.get("reports_ready", False),
        "completion_time": task.get("completion_time"),
        "timestamp": datetime.now(KW).isoformat()
    }
    
    print(f"📊 Status check for task {task_id}: {response['status']}, reports: {response['reports_available']}")
    
    return response

if __name__ == "__main__":
    import os
    
    # Cloud deployment configuration
    port = int(os.getenv("PORT", 8000))
    host = "0.0.0.0"  # Required for Render
    
    print(f"🚀 Starting Kuwait Banking News Analyzer API v2.2 - CLOUD DEPLOYMENT...")
    print(f"🔧 AsyncIO Loop Compatibility: FIXED")
    print(f"🔗 Server will run on: http://{host}:{port}")
    print("📚 API Documentation available at /docs")
    print(f"🌍 Environment: {'PRODUCTION' if os.getenv('RENDER') else 'DEVELOPMENT'}")
    print("🎯 Methodology: Master Orchestrator Integration - FIXED")
    print("📄 Frontend Sync: Enhanced Real-time Updates - FIXED")
    print("📊 Report Generation: Content Transfer Issues - FIXED")
    print("🔌 WebSocket Stability: Connection Issues - FIXED")
    print("📊 Available Banks:", len(BANK_CONFIGS))
    print("📰 Available Newspapers:", len(NEWSPAPER_SCRIPTS))
        
    # Use default event loop policy to avoid uvloop conflicts
    asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())
    
    uvicorn.run(
        "main:app",
        host=host,
        port=port,
        reload=False,  # Disable reload in production
        log_level="info",
        loop="asyncio"
    )